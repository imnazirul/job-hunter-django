"""Get plain text out of a PDF or DOCX upload.

Scanned CVs are the common failure here: a PDF that is really a photo yields no
text at all. We detect that and say so instead of handing an empty string to the
LLM and getting a confidently empty profile back.
"""

import logging
import re

logger = logging.getLogger(__name__)

PDF_MAGIC = b"%PDF"
ZIP_MAGIC = b"PK\x03\x04"
DOC_MAGIC = b"\xd0\xcf\x11\xe0"  # legacy Word binary format

MIN_USEFUL_CHARS = 120


class CVExtractionError(Exception):
    pass


class UnsupportedCV(CVExtractionError):
    pass


class EmptyCVText(CVExtractionError):
    pass


def extract_text(file_obj, filename):
    """Return cleaned text. file_obj must be seekable and open in binary mode."""
    file_obj.seek(0)
    head = file_obj.read(8)
    file_obj.seek(0)

    if head.startswith(DOC_MAGIC):
        raise UnsupportedCV(
            "legacy .doc files are not supported; save as .docx or PDF and upload again"
        )

    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if head.startswith(PDF_MAGIC):
        text = _extract_pdf(file_obj)
    elif head.startswith(ZIP_MAGIC):
        text = _extract_docx(file_obj)
    else:
        raise UnsupportedCV(
            f"file does not look like a PDF or DOCX (extension {suffix or 'none'}, "
            f"first bytes {head[:4]!r})"
        )

    text = tidy(text)
    if len(text) < MIN_USEFUL_CHARS:
        raise EmptyCVText(
            "almost no text could be extracted; if this is a scanned CV, "
            "export a text-based PDF or paste the text manually"
        )
    return text


def _extract_pdf(file_obj):
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(file_obj)
    except PdfReadError as exc:
        raise UnsupportedCV(f"the PDF could not be opened: {exc}") from exc

    if reader.is_encrypted:
        # Some CVs are encrypted with an empty owner password, which decrypts fine.
        try:
            if not reader.decrypt(""):
                raise UnsupportedCV("the PDF is password protected")
        except (PdfReadError, NotImplementedError) as exc:
            raise UnsupportedCV(f"the PDF is password protected: {exc}") from exc

    parts = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            parts.append(page.extract_text() or "")
        except (PdfReadError, ValueError, KeyError) as exc:
            # One broken page should not lose the other five.
            logger.warning("skipped unreadable PDF page %s: %s", number, exc)
            parts.append("")
    return "\n".join(parts)


def _extract_docx(file_obj):
    import zipfile

    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = Document(file_obj)
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        raise UnsupportedCV("the file is a zip archive but not a readable .docx document") from exc
    except KeyError as exc:
        # python-docx reads [Content_Types].xml straight out of the zip, so a zip
        # that is not an Office document surfaces as a KeyError from zipfile.
        raise UnsupportedCV(f"the .docx package is missing a required part: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    # Plenty of CVs lay everything out in tables, so skipping them loses the CV.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cell for cell in cells if cell))
    return "\n".join(parts)


def tidy(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\xa0", " ").replace("\u200b", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
