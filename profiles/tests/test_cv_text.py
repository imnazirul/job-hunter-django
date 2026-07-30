import io

import pytest

from profiles.cv_text import EmptyCVText, UnsupportedCV, extract_text


def docx_bytes(paragraphs, table_rows=None):
    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for cell_index, value in enumerate(row):
                table.cell(row_index, cell_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


LONG_PARAGRAPH = (
    "Senior Backend Engineer with eight years of experience building Django and "
    "PostgreSQL services for high traffic products across several industries."
)


def test_reads_a_docx():
    text = extract_text(docx_bytes([LONG_PARAGRAPH, "Python, Django, PostgreSQL"]), "cv.docx")
    assert "Senior Backend Engineer" in text
    assert "PostgreSQL" in text


def test_reads_docx_tables():
    text = extract_text(
        docx_bytes([LONG_PARAGRAPH], table_rows=[["Skill", "Years"], ["Python", "8"]]),
        "cv.docx",
    )
    assert "Skill | Years" in text
    assert "Python | 8" in text


def test_docx_with_almost_no_text_is_rejected():
    with pytest.raises(EmptyCVText):
        extract_text(docx_bytes(["Alice"]), "cv.docx")


def test_legacy_doc_is_rejected_with_a_useful_message():
    legacy = io.BytesIO(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 500)
    with pytest.raises(UnsupportedCV, match="legacy"):
        extract_text(legacy, "cv.doc")


def test_a_png_pretending_to_be_a_pdf_is_rejected():
    fake = io.BytesIO(b"\x89PNG\r\n\x1a\n" + b"\x00" * 200)
    with pytest.raises(UnsupportedCV):
        extract_text(fake, "cv.pdf")


def test_corrupt_pdf_is_rejected():
    broken = io.BytesIO(b"%PDF-1.7\nthis is not really a pdf body")
    with pytest.raises(UnsupportedCV):
        extract_text(broken, "cv.pdf")


def test_empty_file_is_rejected():
    with pytest.raises(UnsupportedCV):
        extract_text(io.BytesIO(b""), "cv.pdf")


def test_zip_that_is_not_a_docx_is_rejected():
    import zipfile

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("hello.txt", "not a word document")
    buffer.seek(0)
    with pytest.raises(UnsupportedCV):
        extract_text(buffer, "cv.docx")
